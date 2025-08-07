from flask import Flask, request, send_file, jsonify, render_template_string
from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import barcode
from barcode.writer import ImageWriter
from PIL import Image


app = Flask(__name__)

# HTML frontend
HTML = """
<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <title>Δημιουργία Barcode</title>
    <style>
        body {
            font-family: sans-serif;
            max-width: 800px;
            margin: auto;
            padding: 20px;
            background-image: url('/static/background.jpg');
            background-size: cover;
            background-repeat: repeat;
        }
        #logo {
            position: fixed;
            top: 10px;
            left: 10px;
            width: 200px;
            height: auto;
            z-index: 1000;
        }
        button {
            background-color: #007BFF; 
            color: white;              
            border: none;
            padding: 10px 15px;
            cursor: pointer;
            border-radius: 4px;
            font-weight: bold;
            transition: background-color 0.3s ease;
        }

        button:hover {
            background-color: #0056b3; /* πιο σκούρο μπλε στο hover */
        }

        input { margin: 5px 0; width: 100%; padding: 8px; }
        table { width: 100%; border-collapse: collapse; margin-top: 20px; background-color: #f5f5f5; border: 1px solid #ccc;  }
        th, td { border: 1px solid #ccc; padding: 8px; text-align: left; background-color: white;}
        button { padding: 10px 15px; margin-top: 10px; }
        td > button { margin-right: 5px; }
        h2 { text-align: center; }
    </style>
</head>
<body>
<img src="/static/logo.png" alt="Logo" id="logo" />
<h2>Καταχώριση Προϊόντων</h2>
<form id=\"productForm\">
    <input type=\"text\" id=\"barcode\" placeholder=\"Barcode\" required><br>
    <input type=\"text\" id=\"description\" placeholder=\"Περιγραφή\" required><br>
    <input type=\"text\" id=\"code\" placeholder=\"7ψήφιος Κωδικός SAP\" maxlength=\"7\" required><br>
    <button type=\"submit\">Προσθήκη</button>
</form>
<table id=\"productsTable\">
    <thead>
        <tr><th>Barcode</th><th>Περιγραφή</th><th>Κωδικός SAP</th><th>Ενέργειες</th></tr>
    </thead>
    <tbody></tbody>
</table>
<button onclick=\"downloadDoc()\">Κατεβάστε .doc</button>
<script>
    const form = document.getElementById('productForm');
    const table = document.getElementById('productsTable').querySelector('tbody');
    const products = [];
    let editIndex = -1;

    form.onsubmit = function(e) {
        e.preventDefault();
        const barcode = document.getElementById('barcode').value;
        const description = document.getElementById('description').value;
        const code = document.getElementById('code').value;

        if (editIndex === -1) {
            products.push({ barcode, description, code });
        } else {
            products[editIndex] = { barcode, description, code };
            editIndex = -1;
        }

        updateTable();
        form.reset();
    };

    function updateTable() {
        table.innerHTML = '';
        products.forEach((item, index) => {
            const row = document.createElement('tr');
            row.innerHTML = `
                <td>${item.barcode}</td>
                <td>${item.description}</td>
                <td>${item.code}</td>
                <td>
                    <button onclick=\"editProduct(${index})\">✏️</button>
                    <button onclick=\"deleteProduct(${index})\">🗑️</button>
                </td>`;
            table.appendChild(row);
        });
    }

    function editProduct(index) {
        const product = products[index];
        document.getElementById('barcode').value = product.barcode;
        document.getElementById('description').value = product.description;
        document.getElementById('code').value = product.code;
        editIndex = index;
    }

    function deleteProduct(index) {
        products.splice(index, 1);
        updateTable();
    }

    function downloadDoc() {
        fetch('/generate_doc', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ products })
        })
        .then(response => response.blob())
        .then(blob => {
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = 'products.docx';
            a.click();
        });
    }
</script>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(HTML)

@app.route('/generate_doc', methods=['POST'])
def generate_doc():
    data = request.json
    products = data.get('products', [])

    doc = Document()

    section = doc.sections[-1]
    section.page_height = Mm(150)
    section.page_width = Mm(100)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Mm(10)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.bottom_margin = Mm(10)

    for idx, item in enumerate(products):
        if idx > 0:
            doc.add_page_break()

        # Δημιουργία Barcode και αποθήκευση
        barcode_stream = BytesIO()
        code128 = barcode.get('code128', item['barcode'], writer=ImageWriter())
        code128.write(barcode_stream)
        barcode_stream.seek(0)

        #Αποθήκευση barcode σαν PIL γιατι το python-docx αρνήτε για κάποιο λόγο να την πάρει σαν κάτι άλλο
        img = Image.open(barcode_stream)
        img_buffer = BytesIO()
        img.save(img_buffer, format="PNG")
        img_buffer.seek(0)

        # Προσθήκη εικόνας στο doc
        doc.add_picture(img_buffer, width=Mm(60))

        # Προσθήκη υπόλοιπων 
        doc.add_paragraph(item['description'])
        doc.add_paragraph(item['code'])

        # Κεντράρισμα barcode
        barcode_paragraph = doc.add_paragraph()
        barcode_run = barcode_paragraph.add_run()
        barcode_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        barcode_run.add_picture(img_buffer, width=Mm(60))
        

        # Κεντράρισμα Κειμένων
        desc_para = doc.add_paragraph(item['description'])
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        code_para = doc.add_paragraph(item['code'])
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Αποθήκευσh buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='products.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=True, host='0.0.0.0', port=port)
